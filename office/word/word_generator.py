from docx import Document


class WordGenerator:

    async def generate(self, text, output):

        doc = Document()

        doc.add_heading("MARVIS REPORT", level=1)

        doc.add_paragraph(text)

        doc.save(output)

        return output
